from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

# create document
document = Document()

# Picture
document.add_picture(
    'profil.jpg', width=Inches(1.0)
    )

# name,  phone number and email details
speak('You are welcome to the cv generator, ')
speak('What is your name ? ')
name = input('What is your name ? ')
speak('Hello' + name)
speak('What is your Phone Number ? ')
phone_number = input('What is your Phone Number ? ')
speak('What is your E-mail ? ')
email = input('What is your E-mail ? ')
document.add_paragraph(
    name + ' | ' + phone_number +  ' | ' +  email
)

# About me
document.add_heading('About Me')
speak('Tell me about yourself ? ')
about_me = input('Tell me about yourself ? ')
document.add_paragraph(about_me)

#work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter Company ')
from_date = input('From date ')
to_date = input('To date ')

p.add_run(company + ' ').bold= True
p.add_run(from_date + '-' + to_date + '\n').italic= True

experience_details = input(
    'Describe your experience at ' + company)
p.add_run(experience_details)

# More experiences
while True:
    has_more_experiences = input('Do you have more experience (Yes or No) ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter Company ')
        from_date = input('From date ')
        to_date = input('To date ')

        p.add_run(company + ' ').bold= True
        p.add_run(from_date + '-' + to_date + '\n').italic= True

        experience_details = input(
            'Describe your experience at ' + company )
        p.add_run(experience_details)
    else:
        break

# skills
document.add_heading('Skills')
skill = input('Enter skill ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more skills (Yes or No) ')
    if has_more_skills == 'yes':
        skill = input('Enter skill ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

#Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using rekiding's code "



#save document
document.save('cv.docx')